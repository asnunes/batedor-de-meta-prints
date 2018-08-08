# coding=UTF-8
from __future__ import division

import ConfigParser
import sys
import PIL.Image
import cairo
import gi
import glob

import os
import pickle
import pyPdf
#import pytesseract
import re
import shutil
import string
import tempfile

gi.require_version('Gtk', '3.0')
gi.require_version('Notify', '0.7')
from gi.repository import Gtk, Gdk, Notify, GdkPixbuf, GObject
from threading import Thread
from wand.image import Image
from docx import Document
#from docx.shared import Inches


class ABCD(Thread):
    def __init__(self, application):
        super(ABCD, self).__init__()

        self.application = application
        self.ispdf = application.ispdf
        self.loadpath = application.filepath
        self.savepath = application.path_img
        self.page = application.docproperties[0]

    def run(self):

        self.application.loadbar.set_text('Carregando arquivo...')
        self.load_file()
        self.load_coords()
        self.application.loadbar.set_text('Carregando arquivo...')

    def load_file(self):
        if self.ispdf:
            if not os.path.isfile(self.loadpath.decode('utf8')) or (self.application.docproperties[1] != -1
                                                     and pyPdf.PdfFileReader(open(self.loadpath.decode('utf8'), 'rb')).getNumPages() != self.application.docproperties[1]):
                self.application.filepath = ''

                self.loadpath = ''
                return

            with Image(filename=self.loadpath + "[" + str(self.page) + "]", resolution=300) as img:
                img.save(filename=self.savepath)

        else:
            if self.application.docproperties[2]:
                if not os.path.isfile(self.loadpath):
                    self.application.filepath = ''
                    self.loadpath = ''
                    return

                filetoload = PIL.Image.open(self.loadpath)
            else:
                if not os.path.isfile(self.loadpath.decode('utf8')):
                    self.application.filepath = ''
                    self.loadpath = ''
                    return

                filetoload = PIL.Image.open(self.loadpath.decode('utf8'))

            filetoload.save(self.savepath)

    # Carrega os coords
    def load_coords(self):
        count = len(self.application.pcoords)
        loops = 0

        while count > 0:
            # Carrega lista de coordenas
            pcoord = self.application.pcoords[loops]
            x0 = pcoord.x0
            y0 = pcoord.y0
            xf = pcoord.xf
            yf = pcoord.yf
            page = pcoord.page
            primalprintfile = pcoord.primalprintfile
            count -= 1

            if page != self.application.docproperties[0]:
                self.application.loadbar.set_text('Processando página ' + str(page+1) + ' de ' + str(self.application.docproperties[1]))

            coord = Coord(x0, y0, xf, yf, page)
            self.application.coords.append(coord)

            if self.ispdf and (self.page != page or loops == 0):
                self.page = page
                self.load_file()

                img = PIL.Image.open(self.application.path_img)

                # Verificar isso
                img.thumbnail((int(img.size[0] * 0.5), int(img.size[1] * 0.5)), PIL.Image.ANTIALIAS)

                self.application.image = img
            elif self.page != page or loops == 0:
                self.loadpath = self.application.docproperties[2][page]
                self.load_file()

                img = PIL.Image.open(self.application.path_img)

                # Verificar isso
                img.thumbnail((int(img.size[0] * 0.5), int(img.size[1] * 0.5)), PIL.Image.ANTIALIAS)

                self.application.image = img

            coord.generate_pil(self.application)

            if primalprintfile is not None:
                name = primalprintfile.name
                print_type = primalprintfile.thistype
                number = primalprintfile.number
                image = coord.pil
                printfile = PrintFile(image, name, print_type, number)
                coord.set_printfile(printfile)
                self.application.cCoord = coord

            if self.ispdf and (self.application.docproperties[0] != page or loops == 0):
                self.application.docproperties[0] = float(page)
                # Seta o seletor de páginas
                # value, minValue, maxValeu, 1,10,0
                adjustment = Gtk.Adjustment(float(page+1), 1, self.application.docproperties[1], 1, 10, 0)
                self.application.pagespin.set_adjustment(adjustment)

            elif self.application.docproperties[0] != page or loops == 0:
                self.application.docproperties[0] = page
                self.application.pngpicker.set_active(page)

            loops += 1

        self.application.pcoords = []

class Application:

    LEFT_BUTTON = 1
    RIGHT_BUTTON = 3

    def __init__(self):
        # Checa a presença dos arquivos essenciais:
        if self.check_file('bmBase.glade'):
            return
        if self.check_file('default.docx'):
            return
        if self.check_file('about.glade'):
            return
        if self.check_file('default-src.docx'):
            return
        if self.check_file('deleteDialog.glade'):
            return
        if self.check_file('icon.png'):
            return
        if self.check_file('docdialog.glade'):
            return
        if self.check_file('namerMenu.glade'):
            return
        if self.check_file('opendialog.glade'):
            return
        if self.check_file('settings.glade'):
            return

        sys.stdout = open('log.txt', 'wb')

        self.abcd_thread = None

        # Variáveis
        self.count = 0
        self.inicialCoord = [0,0]
        self.finalCoord = [0,0]
        self.coords = []
        self.filepath = ''
        self.path_img = ''
        self.image = None
        self.isleftpressed = False
        self.isrightpressed = False
        self.timeout_id = None
        self.isload = False

        #verifica se está tudo ok com o docx:
        self.isdocxok = True

        # Lista de coord salvas, até alguem carregar é vazia
        self.pcoords = []

        # Verifica se já chamou uma resposta para mudar o nome para o primeiro:
        self.calledanswer = False

        # Carrega .ini e seta valores:
        self.config = ConfigParser.ConfigParser()
        self.listofsetvariables = self.load_ini()

        # pdf ou imagem?
        self.ispdf = False

        # Atributos de pdf 'pagina atual','numero total de paginas', 'lista de caminho, pdf só tem um'
        self.docproperties = [0,-1, []]

        # Adiciona o modelo
        self.listmodel = Gtk.ListStore(str, int, int, int, str)

        # Progress Bar
        self.statusbar = None

        # Coordenada selecionada
        self.cCoord = None

        # Pixbuf base
        self.pixbuf = GdkPixbuf.Pixbuf.new_from_file('./base.png')
        # Pixbuf cortado
        self.pixbuf2 = GdkPixbuf.Pixbuf.new_from_file('./base.png')

        # Inicia notificação
        Notify.init("Batedor de Meta")
        self.print_created_notification = Notify.Notification.new('Print Criado!')

        # Diretório padrão
        self.tmpdir = tempfile.mkdtemp()

        # Diretório com os prints
        self.tmpprintdir = tempfile.mkdtemp()

        # Builder do glade
        self.builder = Gtk.Builder()
        self.builder.add_from_file('bmBase.glade')
        #self.builder.add_from_file('pdfMenu.glade')
        self.builder.add_from_file('namerMenu.glade')

        # Pega os sinais e os conecta
        handlers = {
            "on_entry_changed": self.on_entry_changed,
            "on_value_changed": self.on_value_changed,
            "on_key_pressed": self.on_key_pressed,
        }
        self.builder.connect_signals(handlers)

        # Widgets usados
        self.area = self.builder.get_object('area')
        self.drawScroll = self.builder.get_object('drawScroll')
        self.aboutButton = self.builder.get_object('aboutButton')
        self.namerLayout = self.builder.get_object('namerLayout')
        self.printImage = self.builder.get_object('printImage')
        self.imageScroll = self.builder.get_object('imageScroll')
        self.nameEntry = self.builder.get_object('nameEntry')
        self.backButton = self.builder.get_object('backButton')
        self.okButton = self.builder.get_object('okButton')
        self.deleteButton = self.builder.get_object('deleteButton')
        self.countLabel = self.builder.get_object('countLabel')
        self.saveButton = self.builder.get_object('saveButton')
        self.loadbar = self.builder.get_object('loadbar')

        #   Toolbar
        self.pagespin = self.builder.get_object('pagespin')
        self.pagelabel = self.builder.get_object('pagelabel')
        self.pngpicker = self.builder.get_object('pngpicker')
        self.settingsButton = self.builder.get_object('settingsButton')
        self.pngpicker.connect('changed', self.on_png_picker_changed)
        self.settingsButton.connect('clicked', self.on_settings_button_clicked)

        # Namerlayout
        self.addbutton = self.builder.get_object('addbutton')
        self.deductbutton = self.builder.get_object('deductbutton')

        # RadioButtons
        self.radiobutton1 = self.builder.get_object('radiobutton1')
        self.radiobutton2 = self.builder.get_object('radiobutton2')
        self.radiobutton3 = self.builder.get_object('radiobutton3')

        # Eventos
        self.radiobutton1.connect("toggled", self.on_button_toggled, "1")
        self.radiobutton2.connect("toggled", self.on_button_toggled, "2")
        self.radiobutton3.connect("toggled", self.on_button_toggled, "3")

        # Janela principal
        self.window = self.builder.get_object('batedorDeMeta')
        self.css()
        self.mainLayout = self.builder.get_object('mainLayout')

        # Configura botão de abrir
        self.openButton = self.builder.get_object('openButton')
        self.openButton.connect('clicked', self.open_button_clicked)

        # Configura botão de salvar
        self.saveButton.connect('clicked', self.save_button_clicked)

        # Configura botão sobre
        self.aboutButton.connect('clicked', self.about_button_clicked)

        # Habilita o draw para receber eventos de mudança
        self.area.connect("draw", self.expose)
        self.area.set_events(Gdk.EventMask.BUTTON_PRESS_MASK)
        self.area.add_events(Gdk.EventMask.BUTTON_RELEASE_MASK)
        self.area.add_events(Gdk.EventMask.POINTER_MOTION_MASK)
        self.area.add_events(Gdk.EventMask.ENTER_NOTIFY_MASK)
        self.area.add_events(Gdk.EventMask.LEAVE_NOTIFY_MASK)
        self.area.add_events(Gdk.EventMask.SCROLL_MASK)
        self.area.connect("button-press-event", self.area_event)
        self.area.connect("button-release-event", self.area_event)
        self.area.connect("motion-notify-event", self.area_event)
        self.area.connect("enter-notify-event", self.area_event)
        self.area.connect("leave-notify-event", self.area_event)
        self.area.connect("scroll-event", self.area_event)

        self.window.get_root_window().set_cursor(Gdk.Cursor(Gdk.CursorType.ARROW))
        self.window.set_size_request(1500, 700)

        # Cria o nomeador
        self.mainLayout.add(self.namerLayout)
        self.mainLayout.reorder_child(self.namerLayout, 3)

        # Configura o botão de OK
        self.okButton.connect('clicked', self.ok_button_clicked)

        # Configura o botão de Voltar
        self.backButton.connect('clicked', self.back_button_clicked)

        # Configura o botão de deletar
        self.deleteButton.connect('clicked', self.delete_button_clicked)

        # Configura o botão de adicionar ou diminuir
        self.addbutton.connect('clicked', self.add_number)
        self.deductbutton.connect('clicked', self.add_number)

        # Carrega o previewer
        self.previewer = self.builder.get_object('previewer')
        self.previewer.connect("draw", self.previewer_draw)

        # Prepara a janela
        self.window.connect('delete-event', Gtk.main_quit)
        self.window.set_title("Batedor de Meta")

        try:
            self.window.set_default_icon_from_file('dexter_impressionado.png')
        except:
            pass

        self.window.show_all()

        self.addbutton.get_visible()

        # Deixa os não defaut como invisível
        self.drawScroll.set_visible(False)
        self.namerLayout.set_visible(False)
        self.previewer.set_visible(False)
        self.loadbar.set_visible(False)
        self.pagespin.set_visible(False)
        self.pagelabel.set_visible(False)
        self.pngpicker.set_visible(False)

    def check_file(self, file):
        bool = not os.path.isfile(file)

        if bool:
            dialog = Gtk.MessageDialog(None, 0, Gtk.MessageType.ERROR,
                                       Gtk.ButtonsType.CLOSE, "O arquivo essencial " + file + ' não foi encontrado.')
            dialog.format_secondary_text(
                "O programa irá fechar...")

            dialog.run()
            Gtk.main_quit()
            dialog.destroy()

        return bool

    def css(self):
        self.window.set_name('BM')

        style_provider = Gtk.CssProvider()

        css = """
        #BM {
            background-color: #FFF;
        }

        #BM GtkToolbar {
            background: #FFF;
        }

        #Bm GtkProgressBar{
            font-size: 5em
        }

        """

        style_provider.load_from_data(css)

        Gtk.StyleContext.add_provider_for_screen(
            Gdk.Screen.get_default(),
            style_provider,
            Gtk.STYLE_PROVIDER_PRIORITY_APPLICATION
        )

    def on_key_pressed(self, widget, e, data=None):
        if e.keyval == Gdk.KEY_Delete and self.drawScroll.get_visible() and self.cCoord is not None:
            self.delete_button_clicked(self.deleteButton)
        if e.keyval == Gdk.KEY_Return and self.drawScroll.get_visible() and self.okButton.get_sensitive():
            self.ok_button_clicked(self.okButton)

    def open_button_clicked(self, widget):
        # Faz a limpa
        if self.filepath != '':
            dialog = Gtk.MessageDialog(self.window, 0, Gtk.MessageType.QUESTION,
                                       Gtk.ButtonsType.YES_NO, "Deseja abrir um novo arquivo? Todo o progesso não salvo será perdido")
            dialog.format_secondary_text(
                "Esta operação é irreversível.")

            response = dialog.run()
            if response == Gtk.ResponseType.YES:
                dialog.destroy()
                self.abcd_thread = None
                self.count = 0
                self.coords = []
                self.filepath = ''
                self.path_img = ''
                self.image = None
                self.isleftpressed = False
                self.isrightpressed = False
                self.pcoords = []
                self.calledanswer = False
                self.ispdf = False
                self.docproperties = [0, -1, []]
                self.cCoord = None
                self.tmpdir = tempfile.mkdtemp()
                self.tmpprintdir = tempfile.mkdtemp()
                self.saveButton.set_sensitive(False)
                self.drawScroll.set_visible(False)
                self.namerLayout.set_visible(False)
                self.previewer.set_visible(False)
                self.loadbar.set_visible(False)
                self.pagespin.set_visible(False)
                self.pagelabel.set_visible(False)
                self.pngpicker.set_visible(False)

            else:
                dialog.destroy()
                return

        dialog = Gtk.FileChooserDialog("Por favor, escolha um documento ou imagem", self.window,
                                       Gtk.FileChooserAction.OPEN,
                                       (Gtk.STOCK_CANCEL, Gtk.ResponseType.CANCEL,
                                        Gtk.STOCK_OPEN, Gtk.ResponseType.OK))

        self.add_filters(dialog)

        response = dialog.run()
        if response == Gtk.ResponseType.OK:

            self.filepath = os.path.normpath(dialog.get_filename())
            dialog.destroy()

            while Gtk.events_pending():
                Gtk.main_iteration_do(False)

            if self.filepath[-3:] == 'pdf':
                # Carrega o PDF
                self.ispdf = True
                self.new_temp()
            elif self.filepath[-3:] == 'png' or self.filepath[-3:] == 'jpg':
                # Carrega o PNG
                self.ispdf = False
                self.new_temp()
            elif self.filepath[-2:] == 'bm':
                self.isload = True
                # Carrega o projeto
                self.load_file()

        dialog.destroy()

    # Salva o projeto
    def save_button_clicked(self, widget, tfilename=''):

        filename = ''

        if tfilename != '':
            dialog = Gtk.FileChooserDialog("Selecione a pasta e escolha um nome para o projeto", self.window,
                                           Gtk.FileChooserAction.SAVE,
                                           (Gtk.STOCK_CANCEL, Gtk.ResponseType.CANCEL,
                                            "Selecionar", Gtk.ResponseType.APPLY))
            dialog.set_default_size(800, 400)
            dialog.set_current_name('semnome.bm')

            self.add_filters_save(dialog)

            response = dialog.run()
            if response == Gtk.ResponseType.APPLY:
                filename = os.path.normpath(dialog.get_filename())
                dialog.destroy()

            elif response == Gtk.ResponseType.CANCEL:
                dialog.destroy()
                return
        else:
            filename = os.path.normpath(tfilename + '/projeto.bm')

        with open(filename.decode('utf8'), 'wb') as output:
            # Cria a lista coordenas primitivas
            pcoords = []

            for coord in self.coords:
                x0 = coord.x0
                y0 = coord.y0
                xf = coord.xf
                yf = coord.yf
                page = coord.page

                if coord.printFile is not None:
                    pcoord = PrimalCoord(x0, y0, xf, yf, page, coord.printFile.get_primal())
                else:
                    pcoord = PrimalCoord(x0, y0, xf, yf, page)

                pcoords.append(pcoord)

            pickle.dump(pcoords, output, pickle.HIGHEST_PROTOCOL)
            pickle.dump(self.ispdf, output, pickle.HIGHEST_PROTOCOL)
            pickle.dump(self.docproperties, output, pickle.HIGHEST_PROTOCOL)

    def about_button_clicked(self, widget):
        builder = Gtk.Builder()
        builder.add_from_file('about.glade')
        dialog = builder.get_object('aboutdialog')
        dialog.run()

        dialog.destroy()

    def ok_button_clicked(self, widget):

        # Pega os parametros definidos
        image = self.cCoord.get_pil()
        name = self.nameEntry.get_text()
        # Depriciado number = self.nOfPrintsSpin.get_value_as_int()

        if self.radiobutton1.get_active():
            print_type = "E"
        elif self.radiobutton2.get_active():
            print_type = "L"
        else:
            print_type = "R"

        number = self.get_recurrence(name, print_type)

        printfile = PrintFile(image, name, print_type, number)
        self.cCoord.set_printfile(printfile)

        # Configura o nome dos enunciados:
        if print_type == "L":
            index = self.coords.index(self.cCoord)
            count = 1
            cnumber = 101
            while index - count >= 0:
                cprintfile = self.coords[index - count].printFile
                if cprintfile is not None and cprintfile.thistype  == 'E' and cprintfile.number < cnumber:
                    cprintfile.name = name
                    cnumber = cprintfile.number
                    count += 1
                else:
                    break

        # Configura o backbutton:
        if not self.backButton.get_sensitive() and self.coords.index(self.cCoord) > 0:
            self.backButton.set_sensitive(True)

        # muda o cCoord para o próximo se houver:
        changed = False

        for coord in self.coords:
            if coord.printFile is None:
                self.cCoord = coord
                changed = True
                break

        if self.cCoord.page != self.docproperties[0]:
            self.namerLayout.set_visible(False)
            if self.ispdf:
                # Seta o seletor de páginas e chama a mudança em consequência
                self.pagespin.set_value(float(self.cCoord.page) + 1)
            else:
                self.pngpicker.set_active(self.cCoord.page)

        elif changed is True:
            self.show_maker_layout()
        else:
            self.generate_docx()
            # Gera os docxs

    def back_button_clicked(self, widget):
        # Não precisa verificar o index, pois já estará desativado se não for o correto
        index = self.coords.index(self.cCoord)

        try:
            self.cCoord = self.coords[index - 1]

            if self.cCoord.page != self.docproperties[0]:
                self.namerLayout.set_visible(False)
                if self.ispdf:
                    # Seta o seletor de páginas e chama a mudança em consequência
                    self.pagespin.set_value(float(self.cCoord.page) + 1)
                else:
                    self.pngpicker.set_active(self.cCoord.page)

            else:
                self.show_maker_layout()
                self.okButton.set_label('Próximo não-nomeado')

        except IndexError:
            return

    def delete_button_clicked(self, widget):
        dialog = Gtk.MessageDialog(self.window, 0, Gtk.MessageType.QUESTION,
            Gtk.ButtonsType.YES_NO, "Deseja excluir esse print?")
        dialog.format_secondary_text(
            "Esta operação é irreversível.")

        response = dialog.run()
        if response == Gtk.ResponseType.YES:
            index = self.coords.index(self.cCoord)
            self.coords.remove(self.cCoord)

            if index < len(self.coords):
                self.cCoord = self.coords[index]
            elif index - 1 > 0:
                self.cCoord = self.coords[index - 1]
            else:
                self.cCoord = None
                self.saveButton.set_sensitive(False)

            if self.cCoord is not None and self.cCoord.page != self.docproperties[0]:
                self.namerLayout.set_visible(False)
                if self.ispdf:
                    # Seta o seletor de páginas e chama a mudança em consequência
                    self.pagespin.set_value(float(self.cCoord.page) + 1)
                else:
                    self.pngpicker.set_active(self.cCoord.page)

            else:
                self.show_maker_layout()

        dialog.destroy()

    def on_button_toggled(self, widget, thisid):
        name = self.nameEntry.get_text()
        new_name = name

        # Se for enunciado
        if thisid == '1' and widget.get_active() and self.listofsetvariables[0]:
            only_text = ''
            digit = -1
            if len(name) >= 2 and name[len(name)-2].isdigit() and not name[len(name)-1].isdigit():
                name = name[:-1]
            count = 0
            while count < len(name) and name[-count:].isdigit():
                only_text = name[:-count]
                digit = int(name[-count:])
                count += 1

            if digit != -1:
                new_name = only_text + str(digit +1)
            else:
                new_name = name

            if self.calledanswer:
                self.calledanswer = False

        # Se for item e for com letras
        elif thisid == '2' and widget.get_active():
            if len(name) >= 1 and name[len(name)-1].isdigit():
                only_text = ''
                digit = -1
                count = 0
                while count < len(name) and name[-count:].isdigit():
                    only_text = name[:-count]
                    digit = int(name[-count:])
                    count += 1

                if digit != -1:
                    new_name = only_text + str(digit - 1)
                else:
                    new_name = name

                # Se o item tiver letras (definido por usuario)
                if self.listofsetvariables[0]:
                    new_name += 'a'
            else:
                new_name = name

            if self.calledanswer:
                self.calledanswer = False

        # Se for resposta
        elif thisid == '3' and widget.get_active() and not self.calledanswer and self.coords[0].printFile is not None:
            new_name = self.coords[0].printFile.name
            self.calledanswer = True

        self.nameEntry.set_text(new_name)

    @staticmethod
    def add_filters(dialog):
        filter_only = Gtk.FileFilter()
        filter_only.set_name("Documentos e imagens")
        filter_only.add_pattern("*.pdf")
        filter_only.add_pattern("*.png")
        filter_only.add_pattern("*.jpg")
        dialog.add_filter(filter_only)

        filter_only2 = Gtk.FileFilter()
        filter_only2.set_name("Projetos salvos")
        filter_only2.add_pattern("*.bm")
        dialog.add_filter(filter_only2)

    @staticmethod
    def add_filters_save(dialog):
        filter_only = Gtk.FileFilter()
        filter_only.set_name("BM Projetos")
        filter_only.add_pattern("*.bm")
        dialog.add_filter(filter_only)

    @staticmethod
    def filter_pdf(dialog):
        filter_only = Gtk.FileFilter()
        filter_only.set_name("Documentos PDF")
        filter_only.add_pattern("*.pdf")
        dialog.add_filter(filter_only)

    @staticmethod
    def filter_png(dialog):
        filter_only = Gtk.FileFilter()
        filter_only.set_name("Imagens")
        filter_only.add_pattern("*.png")
        filter_only.add_pattern("*.jpg")
        dialog.add_filter(filter_only)

    def load_file(self):

        bmfile = self.filepath

        if bmfile == '':
            return

        with open( bmfile.decode('utf8'), 'rb') as inputfile:
            pcoords = pickle.load(inputfile)

            ispdf = pickle.load(inputfile)

            docproperties = pickle.load(inputfile)

            # Redefine os parametros
            # Carrega PDF ou PNG
            self.ispdf = ispdf
            self.docproperties = docproperties
            if self.ispdf:
                self.filepath = self.docproperties[2]
            else:
                self.filepath = self.docproperties[2][self.docproperties[0]]
            self.pcoords = pcoords
            self.new_temp()

    def new_temp(self):

        self.loadbar.set_visible(True)
        self.timeout_id = GObject.timeout_add(100, self.on_timeout, None)

        if self.filepath != '' or self.isload:
            self.path_img = os.path.join(self.tmpdir, 'pagetoimg.png')

            #Isso estará no Thread
            self.abcd_thread = ABCD(self)
            self.abcd_thread.start()

        # De tempos em tempos verá se o self.path_img não é nulo passando pelo thread

    def show_img(self):

        # Deixa o drawing area visível
        if not self.drawScroll.get_visible():
            self.drawScroll.set_visible(True)

        # # Bloqueia o abrir
        # self.openButton.set_sensitive(False)

        self.previewer.set_visible(True)

        # Pega o tamanho da imagem para ser o tamanho requisitado
        img = PIL.Image.open(self.path_img)

        # Verificar isso
        try:
            img.thumbnail((int(img.size[0]*0.5),int(img.size[1]*0.5)),PIL.Image.ANTIALIAS)
        except:
            pass

        self.image = img
        self.area.set_size_request(img.size[0], img.size[1])

        img.save(self.path_img)

        # atualiza o drawing area
        self.area.queue_draw()

        self.pixbuf = GdkPixbuf.Pixbuf.new_from_file(self.path_img)

    # Mostra o layout do print namer
    def show_maker_layout(self):

        self.area.queue_draw()

        # Se a lista estiver vazia, não exibe nada
        if not self.coords:
            if self.namerLayout.get_visible():
                self.namerLayout.set_visible(False)
            return
        elif not self.saveButton.get_sensitive():
            self.saveButton.set_sensitive(True)

        # Se a imagem não está lá ou tiver sido alterada, adiciona
        for coord in self.coords:
            if coord.get_pil() is None:
                coord.generate_pil(self)
            elif coord.get_updated():
                coord.generate_pil(self)
                coord.set_updated(False)

        # Deixa visível se não estiver
        if not self.namerLayout.get_visible():
            self.namerLayout.set_visible(True)

        # Adiciona a imagem ao nomeador
        if self.cCoord is None:
            self.cCoord = self.coords[0]
            self.backButton.set_sensitive(False)

        elif self.coords.index(self.cCoord) == 0:
            self.backButton.set_sensitive(False)

        elif not self.backButton.get_sensitive():
            self.backButton.set_sensitive(True)

        # Desabilita item e resposta se for o primeiro
        if len(self.coords) == 1 and self.radiobutton2.get_sensitive():
            self.radiobutton1.set_active(True)
            self.radiobutton2.set_active(False)
            self.radiobutton3.set_active(False)
            self.radiobutton2.set_sensitive(False)
            self.radiobutton3.set_sensitive(False)

        elif not self.radiobutton2.get_sensitive():
            self.radiobutton2.set_sensitive(True)
            self.radiobutton3.set_sensitive(True)


        pil = self.cCoord.get_pil()

        # Adiciona imagem ao widget de imagem
        path = os.path.join(self.tmpdir, 'cprint.png')
        pil.save(path)
        self.printImage.set_from_file(path)
        self.printImage.set_size_request(pil.size[0], pil.size[1])

        # Configura o nomeador
        # value, minValue, maxValeu, 1,10,0
        # Depreciado adjustment = Gtk.Adjustment(1, 1, 99, 1, 10, 0)
        # Depreciado self.nOfPrintsSpin.set_adjustment(adjustment)

        # Configura o contador
        self.set_namer_layout_label_counter()

        self.set_ok_label()

        # Agora tenda regular os nomes se existirem
        if self.cCoord.get_printfile() is None:
            self.make_sujections()
            return

        else:
            printfile = self.cCoord.get_printfile()
            self.nameEntry.set_text(printfile.name)

            if printfile.thistype  == 'E':
                self.radiobutton1.set_active(True)
                self.radiobutton2.set_active(False)
                self.radiobutton3.set_active(False)
            elif printfile.thistype  == 'L':
                self.radiobutton1.set_active(False)
                self.radiobutton2.set_active(True)
                self.radiobutton3.set_active(False)
            else:
                self.radiobutton1.set_active(False)
                self.radiobutton2.set_active(False)
                self.radiobutton3.set_active(True)

                # Depreciado self.nOfPrintsSpin.set_value(printfile.number)

    # Verifica se o entry do arquivo está vazio
    def on_entry_changed(self, widget):
        if self.nameEntry.get_text() == "":
            self.okButton.set_sensitive(False)
            self.addbutton.set_sensitive(False)
            self.deductbutton.set_sensitive(False)

        elif not self.okButton.get_sensitive():
            self.okButton.set_sensitive(True)
            self.addbutton.set_sensitive(True)
            self.deductbutton.set_sensitive(True)

        if ("/" in self.nameEntry.get_text() or "_" in self.nameEntry.get_text()
            or "." in self.nameEntry.get_text()) and self.okButton.get_sensitive():
            self.okButton.set_sensitive(False)

    # Pega mundança de valores no seletor de páginas do pdf
    def on_value_changed(self, widget):

        if self.abcd_thread.isAlive():
            return

        if int(self.pagespin.get_value() -1) == self.docproperties[0]:
            return

        if self.image is None:
            return

        self.docproperties[0] = int(self.pagespin.get_value() -1)

        if self.namerLayout.get_visible():
            self.namerLayout.set_visible(False)
            self.drawScroll.set_visible(False)
        elif self.drawScroll.get_visible():
            self.drawScroll.set_visible(False)

        self.previewer.set_visible(False)
        self.new_temp()

    # Paga mudança no seletor de png
    def on_png_picker_changed(self, widget):

        if self.abcd_thread.isAlive():
            return

        if self.image is None:
            return

        index = self.pngpicker.get_active()

        self.docproperties[0] = index

        self.filepath = self.docproperties[2][self.docproperties[0]]

        if self.namerLayout.get_visible():
            self.namerLayout.set_visible(False)
            self.drawScroll.set_visible(False)
        elif self.drawScroll.get_visible():
            self.drawScroll.set_visible(False)

        self.previewer.set_visible(False)
        self.new_temp()

    # Configura o menu de preferências
    def on_settings_button_clicked(self, widget):
        # Chama o status dialog
        builder = Gtk.Builder()
        builder.add_from_file('settings.glade')
        settingsdialog = builder.get_object('settingsDialog')

        self.listofsetvariables = self.load_ini()
        thislistofvariable = self.listofsetvariables

        # Pega as variáveis
        isalpha = self.listofsetvariables[0]
        issaveprints = self.listofsetvariables[1]
        isblackandwhite = self.listofsetvariables[2]
        isocr = self.listofsetvariables[3]

        # Carrega os widgets:
        isalphabutton = builder.get_object('isalphabutton')
        isnumbutton = builder.get_object('isnumbutton')
        issaveprintsbutton = builder.get_object('issaveprintsbutton')
        isblackandwhitebutton = builder.get_object('isblackandwhitebutton')
        isocrbutton = builder.get_object('isocrbutton')

        # Seta os widgets
        isalphabutton.set_active(isalpha)
        isnumbutton.set_active(not isalpha)
        issaveprintsbutton.set_active(issaveprints)
        isblackandwhitebutton.set_active(isblackandwhite)
        isocrbutton.set_active(isocr)

        # Botões do dialog
        button1 = builder.get_object('button1')
        button2 = builder.get_object('button2')

        thisconfig = self.config

        # O que faz quando o botão é clicado
        def on_button1_clicked(application):
            # Salva o .ini
            cfgfile = open("settings.ini", 'w')
            try:
                thisconfig.set('Sobre os itens', 'isAlpha', isalphabutton.get_active())
            except:
                thisconfig.add_section('Sobre os itens')
                thisconfig.set('Sobre os itens', 'isAlpha', isalphabutton.get_active())
            try:
                thisconfig.set('Set', 'issavePrints', issaveprintsbutton.get_active())
                thisconfig.set('Set', 'isBlackandWhite', isblackandwhitebutton.get_active())
                thisconfig.set('Set', 'isOCR', isocrbutton.get_active())
            except:
                thisconfig.add_section('Set')
                thisconfig.set('Set', 'issavePrints', issaveprintsbutton.get_active())
                thisconfig.set('Set', 'isBlackandWhite', isblackandwhitebutton.get_active())
                thisconfig.set('Set', 'isOCR', isocrbutton.get_active())

            thisconfig.write(cfgfile)
            cfgfile.close()
            settingsdialog.destroy()

            #Atualiza a lista
            thislistofvariable[0] = isalphabutton.get_active()
            thislistofvariable[1] = issaveprintsbutton.get_active()
            thislistofvariable[2] = isblackandwhitebutton.get_active()
            thislistofvariable[3] = isocrbutton.get_active()

        def on_button2_clicked(application):
            settingsdialog.destroy()

        button1.connect('clicked', on_button1_clicked)
        button2.connect('clicked', on_button2_clicked)

        settingsdialog.show()

    # Configura o contador do namer layout
    def set_namer_layout_label_counter(self):
        self.countLabel.set_label(str(self.coords.index(self.cCoord) + 1) + "/" + str(len(self.coords)))

    # Funções de draw
    # Função chamada assim que o drawable é exposto
    def expose(self, wid, cr):

        if self.path_img is not '':
            cr.set_source_rgb(1, 1, 1)
            cr.rectangle(0, 0, self.image.size[0], self.image.size[1])
            cr.fill()

            try:
                ims = cairo.ImageSurface.create_from_png(self.path_img)
            except:
                pass

            cr.set_source_surface(ims, 0, 0)
            cr.paint()

        cr.set_line_width(2)

        # Verifica se está selecionado

        cr.set_source_rgb(0.988, 0.688, 0.24)

        xi = self.inicialCoord[0]
        yi = self.inicialCoord[1]

        xf = self.finalCoord[0]
        yf = self.finalCoord[1]

        #xi, yi, width, height
        cr.rectangle(xi, yi, xf-xi, yf-yi)
        cr.stroke_preserve()

        cr.set_source_rgba(0.988, 0.688, 0.24, 0.55)

        cr.fill()

        for i in self.coords:

            # Se estiver selecionado, muda a cor
            if i is self.cCoord and i.page == self.docproperties[0]:
                cr.set_source_rgb(0.464, 0.804, 0.908)
                cr.rectangle(i.x0, i.y0, i.xf - i.x0, i.yf - i.y0)
                cr.stroke_preserve()
                cr.set_source_rgba(0.464, 0.804, 0.908, 0.55)
                cr.fill()

            elif i.page == self.docproperties[0]:
                cr.set_source_rgb(0.988, 0.688, 0.24)
                cr.rectangle(i.x0, i.y0, i.xf - i.x0, i.yf - i.y0)
                cr.stroke_preserve()
                cr.set_source_rgba(0.988, 0.688, 0.24, 0.55)
                cr.fill()

    # Controla os eventos no DrawingArea
    def area_event(self, w, e):

        if e.type == Gdk.EventType.BUTTON_PRESS \
                and e.button == self.LEFT_BUTTON:
            self.isleftpressed = True
            self.inicialCoord = [e.x, e.y]
            self.finalCoord = [e.x, e.y]

        elif e.type == Gdk.EventType.MOTION_NOTIFY and self.isleftpressed:
            self.finalCoord = [e.x, e.y]
            self.area.queue_draw()

        elif e.type == Gdk.EventType.BUTTON_RELEASE and e.button == self.LEFT_BUTTON:
            # Adiciona a lista
            self.isleftpressed = False

            if self.inicialCoord != self.finalCoord:
                if not self.saveButton.get_sensitive():
                    self.saveButton.set_sensitive(True)

                self.coords.append(
                Coord(self.inicialCoord[0], self.inicialCoord[1], self.finalCoord[0], self.finalCoord[1], self.docproperties[0]))
                self.inicialCoord = [0, 0]
                self.finalCoord = [0, 0]
                self.show_maker_layout()
                #self.print_created_notification.show()

        # Parte que trata da seleção
        elif e.type == Gdk.EventType.BUTTON_PRESS \
                and e.button == self.RIGHT_BUTTON:
            for coord in self.coords[::-1]:

                x0 = coord.x0
                y0 = coord.y0
                xf = coord.xf
                yf = coord.yf
                page = coord.page

                xmin = int(min(x0,xf))
                ymin = int(min(y0,yf))
                xmax = int(max(x0,xf))
                ymax = int(max(y0,yf))

                if xmin < e.x < xmax and ymin < e.y < ymax and float(page) == self.docproperties[0]:
                    self.cCoord = coord
                    self.show_maker_layout()
                    self.okButton.set_label('Próximo não-nomeado')
                    break

        if (e.type == Gdk.EventType.BUTTON_PRESS
                and e.button == self.RIGHT_BUTTON and self.cCoord is not None and self.cCoord.edition != self.cCoord.NONE
            and float(self.cCoord.page) == self.docproperties[0]
            ) \
                or (self.isrightpressed and e.type == Gdk.EventType.MOTION_NOTIFY) \
                or (e.type == Gdk.EventType.BUTTON_RELEASE and e.button == self.RIGHT_BUTTON):
            if self.cCoord.edition == self.cCoord.LEFT:
                if self.cCoord.x0 < self.cCoord.xf and self.cCoord.xf > e.x + 10:
                    self.cCoord.x0 = e.x
                elif self.cCoord.x0 > self.cCoord.xf and self.cCoord.x0 > e.x + 10:
                    self.cCoord.xf = e.x
            elif self.cCoord.edition == self.cCoord.RIGHT:
                if self.cCoord.x0 > self.cCoord.xf and self.cCoord.xf < e.x - 10:
                    self.cCoord.x0 = e.x
                elif self.cCoord.x0 < self.cCoord.xf and self.cCoord.x0 < e.x - 10:
                    self.cCoord.xf = e.x
            elif self.cCoord.edition == self.cCoord.TOP:
                if self.cCoord.y0 < self.cCoord.yf and self.cCoord.yf > e.y + 10:
                    self.cCoord.y0 = e.y
                elif self.cCoord.y0 > self.cCoord.yf and self.cCoord.y0 > e.y + 10:
                    self.cCoord.yf = e.y
            elif self.cCoord.edition == self.cCoord.BOTTOM:
                if self.cCoord.y0 > self.cCoord.yf and self.cCoord.yf < e.y - 10:
                    self.cCoord.y0 = e.y
                elif self.cCoord.y0 < self.cCoord.yf and self.cCoord.y0 < e.y - 10:
                    self.cCoord.yf = e.y
            if e.type == Gdk.EventType.BUTTON_RELEASE and e.button == self.RIGHT_BUTTON:
                self.isrightpressed = False
                self.cCoord.updated = True
                self.show_maker_layout()
                self.okButton.set_label('Próximo não-nomeado')
            elif not self.isrightpressed:
                self.isrightpressed = True

            self.area.queue_draw()

        # Trata do cursor para mudança de tamanho
        elif (self.cCoord is not None) and self.cCoord.page == self.docproperties[0] and e.type == Gdk.EventType.MOTION_NOTIFY:

            if self.count > 5:

                cursor = self.window.get_window().get_cursor()

                xmin = min(self.cCoord.x0, self.cCoord.xf)
                ymin = min(self.cCoord.y0, self.cCoord.yf)
                xmax = max(self.cCoord.x0, self.cCoord.xf)
                ymax = max(self.cCoord.y0, self.cCoord.yf)

                cursor1 = Gdk.Cursor.new_from_name(self.window.get_root_window().get_cursor().get_display(),'e-resize')
                cursor2 = Gdk.Cursor.new_from_name(self.window.get_root_window().get_cursor().get_display(), 'n-resize')

                if xmin - 10 < e.x < xmin + 10 and ymin <= e.y <= ymax and cursor != cursor1:
                    self.window.get_window().set_cursor(cursor1)
                    self.cCoord.edition = self.cCoord.LEFT
                elif xmax - 10 < e.x < xmax + 10 and ymin <= e.y <= ymax and cursor != cursor1:
                    self.window.get_window().set_cursor(cursor1)
                    self.cCoord.edition = self.cCoord.RIGHT
                elif ymin - 10 < e.y < ymin + 10 and xmin <= e.x <= xmax and cursor != cursor2:
                    self.window.get_window().set_cursor(cursor2)
                    self.cCoord.edition = self.cCoord.TOP
                elif ymax - 10 < e.y < ymax + 10 and xmin <= e.x <= xmax and cursor != cursor2:
                    self.window.get_window().set_cursor(cursor2)
                    self.cCoord.edition = self.cCoord.BOTTOM
                elif cursor == cursor1 or cursor == cursor2:
                    self.window.get_window().set_cursor(Gdk.Cursor.new_from_name(
                        self.window.get_window().get_cursor().get_display(),'crosshair'))
                    self.cCoord.edition = self.cCoord.NONE

                self.count = 0
            else:
                self.count += 1

        # Parte que trata da mudança no cursor de seleção
        if e.type == Gdk.EventType.ENTER_NOTIFY:
            cursor = Gdk.Cursor.new_from_name(self.window.get_root_window().get_cursor().get_display(),'crosshair')
            self.window.get_window().set_cursor(cursor)

        elif e.type == Gdk.EventType.LEAVE_NOTIFY:
            cursor = Gdk.Cursor.new_from_name(self.window.get_root_window().get_cursor().get_display(),'default')
            self.window.get_window().set_cursor(cursor)

        # Lida com o previewer
        if e.type == Gdk.EventType.MOTION_NOTIFY or e.type == Gdk.EventType.SCROLL:
            self.set_previewer(e)

    # seta o previewer
    def set_previewer(self, e):
        x0 = int(e.x)
        y0 = int(e.y)

        if self.pixbuf is not None:
            if x0 - 25 > 0 and y0 - 13 > 0 and x0 + 25 < self.pixbuf.get_width() and y0 + 13 < self.pixbuf.get_height():
                pixbuf = self.pixbuf.new_subpixbuf(x0 - 25, y0 - 13, 50, 25)
                self.pixbuf2 = pixbuf.scale_simple(100, 50, GdkPixbuf.InterpType.BILINEAR)

            else:
                self.pixbuf2 = GdkPixbuf.Pixbuf.new_from_file('./base.png')

            self.previewer.queue_draw()

    # Desenha no previewer
    def previewer_draw(self, w, cr):
        GdkPixbuf.Pixbuf.new_from_file('./base.png')
        Gdk.cairo_set_source_pixbuf(cr, self.pixbuf2, 0, 0)
        cr.paint()

        cr.set_source_rgba(0.464, 0.804, 0.908, 0.85)

        # Desenha o cursor
        cr.set_line_width(2)

        cr.set_line_cap(cairo.LINE_CAP_SQUARE)
        cr.move_to(1, 24)
        cr.line_to(46, 24)
        cr.stroke()

        cr.set_line_cap(cairo.LINE_CAP_SQUARE)
        cr.move_to(54, 24)
        cr.line_to(100, 24)
        cr.stroke()

        cr.set_line_cap(cairo.LINE_CAP_SQUARE)
        cr.move_to(50, 1)
        cr.line_to(50, 20)
        cr.stroke()

        cr.set_line_cap(cairo.LINE_CAP_SQUARE)
        cr.move_to(50, 28)
        cr.line_to(50, 52)
        cr.stroke()

    # Define os nomes sugeridos
    def make_sujections(self):
        if self.cCoord is None:
                index = 1
        else:
                index = self.coords.index(self.cCoord)

        name = ''
        thistype = 'E'
        # O quanto deve ser adicionado
        passo = 0

        # Pega um coord de referência
        # Tenta voltar e pegar um
        count = 1
        while index - count >= 0:
            if self.coords[index-count].printFile is not None:
                name = self.coords[index-count].printFile.name
                if self.coords[index - count].printFile is not None:
                    thistype = self.coords[index-count].printFile.thistype
                passo = count
                break
            count += 1

        # Caso não pegue, tenta avançar e pegar um
        if name is '':
            count = 1
            while index + count < len(self.coords):
                if self.coords[index+count].printFile is not None:
                    name = self.coords[index+count].printFile.name
                    if self.coords[index-count].printFile is not None:
                        thistype = self.coords[index - count].printFile.thistype
                    passo = -count
                    break
                count += 1

        # Caso não pegue, sai de cena
        if name is '':
            self.nameEntry.set_text('')
            # Depreciado self.nOfPrintsSpin.set_value(1)
            return

        # Pega os dígitos no número:
        new_only_text = name
        new_only_number = -1
        count2 = 1

        while len(name) - count2 >= 0 and name[len(name) - count2].isdigit():
            new_only_text = name[:-count2]
            new_only_number = int(name[-count2:])
            count2 += 1

        # Se tiver dígito
        if new_only_number != -1:
            new_name = new_only_text + str(new_only_number + passo)
            self.nameEntry.set_text(new_name)

        # Se não tiver dígito
        elif new_only_number == -1 and thistype  == 'L' or (thistype  == 'R' and self.listofsetvariables[0]):
            new_only_text = name[:-1]
            letter = name[len(name) - 1]
            # Pega a letra e seleciona o próximo na lista alfabética:
            letter_index = string.ascii_letters.find(letter)
            if 0 <= letter_index + passo < len(string.ascii_letters):
                letter = string.ascii_letters[letter_index + passo]

            new_name = new_only_text + letter
            self.nameEntry.set_text(new_name)

    # Set label do button
    def set_ok_label(self):
        # Define o botão de ok
        is_done = 0
        for coord in self.coords:
            if coord.printFile is None:
                is_done += 1
            if is_done > 1:
                break

        if is_done > 1 and self.okButton.get_label() is not 'Próximo não-nomeado':
            self.okButton.set_label('Próximo não-nomeado')

        elif is_done <= 1:
            self.okButton.set_label('Gerar Documentos Word')

    # Generate docx
    def generate_docx(self):
        dialog = Gtk.FileChooserDialog("Selecione a pasta onde deseja salvar", self.window,
            Gtk.FileChooserAction.SELECT_FOLDER,
            (Gtk.STOCK_CANCEL, Gtk.ResponseType.CANCEL,
             "Selecionar", Gtk.ResponseType.APPLY))
        dialog.set_default_size(800, 400)

        response = dialog.run()
        if response == Gtk.ResponseType.APPLY:
            filename = dialog.get_filename()
            self.tmpprintdir = tempfile.mkdtemp()

            dialog.destroy()

            # Salva o projeto
            self.save_button_clicked(self.saveButton, filename)

            #Salva os prints
            for coord in self.coords:
                # Verifica se o enunciado não tem item e troca o nome se tiver
                if coord.printFile is not None:
                    coord.printFile.save(self)

            # Chama o status dialog
            dialogbuilder = Gtk.Builder()
            dialogbuilder.add_from_file(os.path.abspath('docdialog.glade'))
            statusdialog = dialogbuilder.get_object('dialog1')

            # Carrega os Widgets
            self.statusbar = dialogbuilder.get_object('progressbar1')
            okdialogbutton = dialogbuilder.get_object('button1')

            # Configura botão de OK
            def on_ok_dialog_button_clicked(widget):
                statusdialog.destroy()

            okdialogbutton.connect('clicked', on_ok_dialog_button_clicked)

            okdialogbutton.set_sensitive(False)

            # Pega o botão e destroi o dialog
            #builder.get_object('okdialogbutton').connect('clicked', Gtk.ResponseType.DELETE_EVENT)

            # Pega a lista
            treeview = dialogbuilder.get_object('treeview7')

            # Adiciona as colunas
            column = Gtk.TreeViewColumn("Questão", Gtk.CellRendererText(), text=0)
            column2 = Gtk.TreeViewColumn("Enunciados", Gtk.CellRendererText(), text=1)
            column3 = Gtk.TreeViewColumn("Itens", Gtk.CellRendererText(), text=2)
            column4 = Gtk.TreeViewColumn("Gabaritos", Gtk.CellRendererText(), text=3)
            column5 = Gtk.TreeViewColumn("Status", Gtk.CellRendererText(), text=4)
            treeview.append_column(column)
            treeview.append_column(column2)
            treeview.append_column(column3)
            treeview.append_column(column4)
            treeview.append_column(column5)

            treeview.set_model(self.listmodel)
            statusdialog.show()

            self.executar(filename)

            # Permite atualiza gui
            while Gtk.events_pending():
                Gtk.main_iteration_do(False)

            # Se for para salver os prints, salva
            if self.listofsetvariables[1]:
                shutil.make_archive(os.path.normpath(filename + '/prints').decode('UTF-8'), format='zip', root_dir=self.tmpprintdir)

            # Permite atualiza gui
            while Gtk.events_pending():
                Gtk.main_iteration_do(False)

            #Atualiza barra
            self.statusbar.set_fraction(1)
            self.statusbar.set_text('Concluído!')

            #Deixa o botão de ok sensível novamente
            okdialogbutton.set_sensitive(True)

        elif response == Gtk.ResponseType.CANCEL:
            dialog.destroy()

    def executar(self, filename):
        lista_de_prints = []

        for coord in self.coords:
            lista_de_prints.append(coord.printFile.file_name())

        lista_de_prints = self.sorted_nicely(lista_de_prints)

        enunciados = []
        perguntas = []
        respostas = []

        for print_path in lista_de_prints:
            index = lista_de_prints.index(print_path)
            index_to_underline = print_path.find('_')

            if print_path[index_to_underline + 1] == 'E':
                # O print eh um enunciado
                # Se a lista enunciado não estiver vazia e as questões forem diferentes, limpa ela

                if enunciados and enunciados[0][:enunciados[0].find('_')] != print_path[:index_to_underline]:
                    enunciados = []
                    perguntas = []
                    respostas = []

                # adiciona esse print a lista de enunciados
                enunciados.append(print_path)

                # Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx
                if index == len(lista_de_prints) - 1 or\
                                print_path[:index_to_underline] != lista_de_prints[index + 1][:lista_de_prints[index + 1].find('_')]:

                    Questao(enunciados, perguntas, respostas, self, filename)
                    self.update_status_bar(index, len(lista_de_prints))

            elif print_path[index_to_underline + 1] == 'L':
                # O print eh uma letra ou item
                # Se a lista perguntas não estiver vazia e as questões forem diferentes, limpa ela
                if perguntas and perguntas[0][:perguntas[0].find('_')] != print_path[0:index_to_underline]:
                    perguntas = []
                    respostas = []

                # adiciona esse print a lista de enunciados
                perguntas.append(print_path)

                # Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx
                if index == len(lista_de_prints) - 1 or print_path[:index_to_underline] != lista_de_prints[index + 1][:lista_de_prints[index + 1].find('_')]:
                    if enunciados:
                        Questao(enunciados, perguntas, respostas, self, filename)
                        self.update_status_bar(index, len(lista_de_prints))


            elif print_path[index_to_underline + 1] == 'R':
                # O print eh uma resposta
                # Se a lista respostas não estiver vazia e as questões forem diferentes, limpa ela
                if respostas and respostas[0][0:respostas[0].find('_')] != print_path[0:index_to_underline]:
                    respostas = []

                # adiciona esse print a lista de enunciados
                respostas.append(print_path)

                # Se a proxima questao tem nome diferente, é porque ela acabou e cria um docx
                if index == len(lista_de_prints) - 1 or print_path[0:index_to_underline] != lista_de_prints[index + 1][0:lista_de_prints[index + 1].find('_')]:
                    if enunciados:
                        Questao(enunciados, perguntas, respostas, self, filename)
                        self.update_status_bar(index, len(lista_de_prints))

    def update_status_bar(self, index, lenght):
        while Gtk.events_pending():
            Gtk.main_iteration_do(False)

        if self.listofsetvariables[1]:
            status = ((index + 1) / (lenght * 1.1))
        else:
            status = ((index + 1) / lenght)

        self.statusbar.set_fraction(status)
        if status == 1:
            self.statusbar.set_text('Concluído!')
        elif status >= 0.85 and self.listofsetvariables[1]:
            self.statusbar.set_text('Gerando arquivos de prints')

    @staticmethod
    def sorted_nicely(l):
        convert = lambda text: int(text) if text.isdigit() else text
        alphanum_key = lambda key: [convert(c) for c in re.split('([0-9]+)', key)]
        return sorted(l, key=alphanum_key)

    #Update progressbar e chama show_img quando o thread está pronto.
    def on_timeout(self, user_data):
        while Gtk.events_pending():
            Gtk.main_iteration_do(False)

        if not self.abcd_thread.isAlive() and self.loadbar.get_visible():
            if self.filepath == '':
                if self.ispdf:
                    dialog = Gtk.FileChooserDialog("PDF não encontrado ou incompatível, por favor escolha outro.",
                                                   self.window,
                                                   Gtk.FileChooserAction.OPEN,
                                                   (Gtk.STOCK_CANCEL, Gtk.ResponseType.CANCEL,
                                                    Gtk.STOCK_OPEN, Gtk.ResponseType.OK))

                    self.filter_pdf(dialog)

                    response = dialog.run()
                    if response == Gtk.ResponseType.OK:
                        self.filepath = os.path.normpath(dialog.get_filename())
                        dialog.destroy()
                        self.new_temp()
                        self.docproperties[2] = self.filepath
                        return
                    else:
                        dialog.destroy()
                        self.loadbar.set_visible(False)
                        return
                else:
                    dialog = Gtk.FileChooserDialog("Imagem não encontrada, especifique outra.",
                                                   self.window,
                                                   Gtk.FileChooserAction.OPEN,
                                                   (Gtk.STOCK_CANCEL, Gtk.ResponseType.CANCEL,
                                                    Gtk.STOCK_OPEN, Gtk.ResponseType.OK))

                    self.filter_png(dialog)

                    response = dialog.run()
                    if response == Gtk.ResponseType.OK:
                        self.filepath = os.path.normpath(dialog.get_filename())
                        dialog.destroy()
                        self.new_temp()
                        self.docproperties[2][self.docproperties[0]] = self.filepath
                        return
                    else:
                        dialog.destroy()
                        self.loadbar.set_visible(False)
                        return

            if self.ispdf:
                # Pega o número de páginas do PDF se já não está definido
                if self.docproperties[1] < 0:
                    reader = pyPdf.PdfFileReader(open(self.filepath.decode('utf8'), 'rb'))
                    self.docproperties = [0, reader.getNumPages(), self.filepath]
                if self.pagelabel.get_label() == 'de 999' or self.isload:
                    # Seta o seletor de páginas
                    # value, minValue, maxValeu, 1,10,0
                    adjustment = Gtk.Adjustment(self.docproperties[0]+1, 1, self.docproperties[1], 1, 10, 0)
                    self.pagespin.set_adjustment(adjustment)
                    self.pagelabel.set_label('de ' + str(self.docproperties[1]))
                    self.isload = False

                self.window.set_title('Batedor de Meta - ' + os.path.basename(self.filepath))

                if not self.pagespin.get_visible():
                    self.pagespin.set_visible(True)
                    self.pagelabel.set_visible(True)

                if self.pngpicker.get_visible():
                    self.pngpicker.set_visible(False)

            else:
                self.window.set_title('Batedor de Meta - ' + os.path.basename(self.filepath))

                if self.docproperties[1] < 0 or self.isload:
                    # Prepara os arquivos de imagem
                    if self.isload:
                        self.filepath = self.filepath.encode('utf8')

                    filefolder = os.path.dirname(self.filepath.decode('utf8'))

                    listofimgs = []

                    for thisfile in glob.glob(filefolder + "\*.png"):
                        listofimgs.append(os.path.normpath(thisfile))
                        print listofimgs.append(os.path.normpath(thisfile))

                    for thisfile in glob.glob(filefolder + "\*.jpg"):
                        listofimgs.append(os.path.normpath(thisfile))

                    listofimgs = self.sorted_nicely(listofimgs)

                    print listofimgs

                    # Cria o docproprities
                    filepath = self.filepath
                    self.docproperties = [listofimgs.index(filepath.decode('utf-8')), len(listofimgs), listofimgs]

                    # Configura o seletor
                    listofimgs_store = Gtk.ListStore(str)
                    for img in listofimgs:
                        listofimgs_store.append([img[len(filefolder):]])

                    self.pngpicker.set_model(listofimgs_store)
                    self.pngpicker.set_active(self.docproperties[0])
                    renderer_text = Gtk.CellRendererText()
                    self.pngpicker.pack_start(renderer_text, True)
                    self.pngpicker.add_attribute(renderer_text, "text", 0)
                    self.isload = False

                if not self.pngpicker.get_visible():
                    self.pngpicker.set_visible(True)

                if self.pagespin.get_visible():
                    self.pagespin.set_visible(False)
                    self.pagelabel.set_visible(False)

            self.show_img()

            if not self.drawScroll.get_visible():
                self.drawScroll.set_visible(True)

            self.show_maker_layout()
            self.loadbar.set_visible(False)

        self.loadbar.pulse()
        return self.loadbar.get_visible()

    def get_recurrence(self, thatname, thattype):

        count = 1

        for coord in self.coords:
            if coord.printFile is not None:
                thisname = coord.printFile.name
                thistype = coord.printFile.thistype

                if thisname == thatname and thistype == thattype:
                    count += 1

        return count

    def add_number(self, widget):

        if widget is self.addbutton:
            more = 1
        else:
            more = -1

        #Adiciona
        name = self.nameEntry.get_text()
        new_only_name = name
        new_only_number = -1
        count = 0

        while count <= len(name) and name[-count:].isdigit():
            new_only_number = int(name[-count:])
            new_only_name = name[:-count]
            count += 1

        # Se for digito:
        if new_only_number != -1:
            if new_only_number > 1 or (new_only_number == 1 and more == 1):
                self.nameEntry.set_text(new_only_name + str(new_only_number + more))
        # Se for letra:
        else:
            new_only_text = name[:-1]
            letter = name[len(name) - 1]
            # Pega a letra e seleciona o próximo na lista alfabética:
            letter_index = string.ascii_letters.find(letter)
            if 0 <= letter_index + more < len(string.ascii_letters):
                letter = string.ascii_letters[letter_index + more]

            new_name = new_only_text + letter
            self.nameEntry.set_text(new_name)

    # Carrega .ini
    def load_ini(self):

        try:
            self.config.read("settings.ini")
        except:
            # Se não encontrou o .ini, cria:
            cfgfile = open("settings.ini", 'w')
            self.config.add_section('Sobre os itens')
            self.config.set('Sobre os itens', 'isAlpha', True)
            self.config.add_section('Set')
            self.config.set('Set', 'issavePrints', False)
            self.config.set('Set', 'isBlackandWhite', False)
            self.config.set('Set', 'isOCR', False)
            self.config.write(cfgfile)
            cfgfile.close()

        # Agora tenta carregar as variáveis:
        try:
            isaplha = self.config.getboolean("Sobre os itens", "isAlpha")
            issaveprints = self.config.getboolean("Set", "issavePrints")
            isblackandwhite = self.config.getboolean("Set", "isBlackandWhite")
            isocr = self.config.getboolean("Set", "isOCR")
        except:
            isaplha = False
            issaveprints = False
            isblackandwhite = False
            isocr = False

        return [isaplha, issaveprints, isblackandwhite, isocr]

class Coord:
    # globais
    NONE = 0
    LEFT = 1
    RIGHT = 2
    TOP = 3
    BOTTOM = 4

    def __init__(self,x0,y0,xf,yf, page, pil=None):

        self.x0 = x0
        self.y0 = y0
        self.xf = xf
        self.yf = yf
        self.page = page
        self.pil = pil
        self.printFile = None

        # Depriciado self.is_selected = False
        self.edition = self.NONE
        self.updated = False

    def generate_pil(self, application):
        # Pega a lista de coordenas minima e máxima

        xi = int(min(self.x0, self.xf))
        yi = int(min(self.y0, self.yf))
        xf = int(max(self.x0, self.xf))
        yf = int(max(self.y0, self.yf))

        # Corta a imagem
        img = application.image.crop((xi, yi, xf, yf))

        # Converte para preto e branco se desejado
        if application.listofsetvariables[2]:
            img = img.convert('L')

        self.set_pil(img)

    def set_pil(self,img):
        self.pil = img

    def get_pil(self):
        return  self.pil

    def set_printfile(self, printfile):
        self.printFile = printfile

    def get_printfile(self):
        return self.printFile

    # Depriciado
    # def set_is_selected(self, is_selected):
    #    self.is_selected = is_selected

    def set_updated(self, updated):
        self.updated = updated

    def get_updated(self):
        return self.updated

class PrintFile:
    def __init__(self, image, name, thistype, number):
        self.image = image
        self.name = name
        self.thistype = thistype
        self.number = number
        self.path = ''

    def file_name(self):
        return self.name + "_" + self.thistype + str(self.number) + ".png"

    def save(self, application):
        path = os.path.join(application.tmpprintdir, self.file_name())
        self.image.save(path)
        self.path = path

    def get_primal(self):
        return PrimalPrintFile(self)

# Classes para salvar objeto
class PrimalCoord:
    def __init__(self, x0, y0, xf, yf, page, primalprintfile=None):
        self.x0 = x0
        self.y0 = y0
        self.xf = xf
        self.yf = yf
        self.page = page
        self.primalprintfile = primalprintfile

class PrimalPrintFile:
    def __init__(self,printfile):
        self.name = printfile.name
        self.thistype = printfile.thistype
        self.number = printfile.number

class Questao:
    def __init__(self, enunciados, itens, respostas, application, savepath):
        self.enunciados = enunciados
        self.itens = itens
        self.respostas = respostas

        self.generate_docx(application, savepath)

    def generate_docx(self, application, savepath):
        isok = application.isdocxok

        if not isok:
            isok = True

        document = Document('default.docx')
        document.add_heading('Enunciado', 0)

        if not self.enunciados:
            isok = False

        self.list_to_docx(self.enunciados, document, application)

        nome = self.enunciados[0][:self.enunciados[0].find('_')]

        if self.itens:
            self.list_to_docx(self.itens, document, application)
            nome = self.itens[0][:self.enunciados[0].find('_')]

        if self.respostas:
            document.add_heading('Resposta', 0)
            self.list_to_docx(self.respostas, document, application)

        document.save(os.path.normpath(savepath + '/' + nome + '.docx').decode('utf8'))

        if isok:
            status = 'Sem enunciado!'
        else:
            status = 'OK'

        application.listmodel.append([nome, len(self.enunciados), len(self.itens), len(self.respostas), status])

    def list_to_docx(self, lista, document, application):
        for path in lista:
            im = PIL.Image.open(os.path.join(application.tmpprintdir, path))

            document.add_picture(os.path.join(application.tmpprintdir, path))

            # Transforma a imagem em texto:
            #if application.listofsetvariables[3]:
                #document.add_paragraph(pytesseract.image_to_string(im).decode('UTF-8'))

            #if imw / 72 < 6:
                #document.add_picture(os.path.join(application.tmpprintdir, path))
            #else:
                #document.add_picture(os.path.join(application.tmpprinaatdir, path), width=Inches(5.5))

if __name__ == "__main__":
    app = Application()
Gtk.main()
